"""Render tailored markdown → .docx → .pdf via LibreOffice headless.

The simplest pipeline that works without modifying the source .docx files:
  1. python-docx writes a fresh .docx from the tailored markdown.
  2. `libreoffice --headless --convert-to pdf` renders that .docx to PDF.
  3. We stamp a UUID into the PDF metadata so per-variant response rate is tracked.

This is intentionally NOT a faithful re-render of your hand-styled .docx
masters. It ships a clean ATS-friendly layout. A future change can revisit
fidelity once we know which variants perform.
"""

from __future__ import annotations

import logging
import re
import shutil
import subprocess
import uuid
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.shared import Pt

from src.config import DATA_DIR

log = logging.getLogger(__name__)

TAILORED_DIR = DATA_DIR / "tailored"


@dataclass
class RenderResult:
    uuid: str
    docx_path: Path
    pdf_path: Path


def render(
    tailored_md: str,
    *,
    variant: str,
    company: str,
    role: str,
) -> RenderResult:
    TAILORED_DIR.mkdir(parents=True, exist_ok=True)
    rid = uuid.uuid4().hex[:12]
    safe_company = re.sub(r"[^A-Za-z0-9_-]+", "_", company)[:32]
    safe_role = re.sub(r"[^A-Za-z0-9_-]+", "_", role)[:48]
    stem = f"{rid}__{variant}__{safe_company}__{safe_role}"

    docx_path = TAILORED_DIR / f"{stem}.docx"
    _md_to_docx(tailored_md, docx_path, uuid_tag=rid)

    pdf_path = _docx_to_pdf(docx_path)

    log.info("rendered tailored resume: %s", pdf_path)
    return RenderResult(uuid=rid, docx_path=docx_path, pdf_path=pdf_path)


def _md_to_docx(md: str, out: Path, *, uuid_tag: str) -> None:
    doc = Document()
    # Set a sane default font for ATS scanners.
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for line in md.splitlines():
        s = line.rstrip()
        if not s.strip():
            doc.add_paragraph("")
        elif s.startswith("# "):
            doc.add_heading(s[2:].strip(), level=0)
        elif s.startswith("## "):
            doc.add_heading(s[3:].strip(), level=1)
        elif s.startswith("### "):
            doc.add_heading(s[4:].strip(), level=2)
        elif s.lstrip().startswith(("- ", "* ")):
            text = s.lstrip()[2:].strip()
            doc.add_paragraph(text, style="List Bullet")
        else:
            doc.add_paragraph(s)

    # Stamp UUID into docx core properties so it propagates to PDF metadata.
    doc.core_properties.identifier = uuid_tag
    doc.core_properties.comments = f"apply-agent uuid={uuid_tag}"
    doc.save(out)


def _docx_to_pdf(docx_path: Path) -> Path:
    libreoffice = shutil.which("libreoffice") or shutil.which("soffice")
    pdf_path = docx_path.with_suffix(".pdf")
    if libreoffice is None:
        log.warning(
            "libreoffice not found — PDF rendering skipped. "
            "Install with: sudo apt install libreoffice --no-install-recommends",
        )
        # Touch a stub so downstream paths exist; dry-run won't actually upload.
        pdf_path.write_bytes(b"")
        return pdf_path

    cmd = [
        libreoffice,
        "--headless",
        "--convert-to", "pdf",
        "--outdir", str(docx_path.parent),
        str(docx_path),
    ]
    try:
        subprocess.run(cmd, check=True, capture_output=True, timeout=60)
    except subprocess.CalledProcessError as e:
        log.error("libreoffice convert failed: %s", e.stderr.decode("utf-8", "replace"))
        raise
    return pdf_path
